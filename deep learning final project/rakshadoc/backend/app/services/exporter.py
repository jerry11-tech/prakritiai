import os
import json
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from app.services.braille_converter import BrailleConverter
from app.services.tts_engine import TTSEngine

class ExporterService:
    @staticmethod
    def export_json(document_dict: dict, output_path: str) -> str:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(document_dict, f, indent=2, ensure_ascii=False)
        return output_path

    @staticmethod
    def export_docx(document, layout_elements, output_path: str) -> str:
        doc = DocxDocument()
        doc.add_heading(f"RakshaDoc Extracted Layout: {document.filename}", level=0)
        
        for elem in layout_elements:
            doc.add_heading(f"[{elem.category}] (Confidence: {elem.confidence*100:.1f}%)", level=2)
            doc.add_paragraph(elem.extracted_text)
            doc.add_paragraph(f"Braille: {elem.braille_text}")
            doc.add_paragraph("------------------------------------")
            
        doc.save(output_path)
        return output_path

    @staticmethod
    def export_pdf(document, layout_elements, output_path: str) -> str:
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, f"RakshaDoc PDF Export: {document.filename}")
        
        y = height - 80
        c.setFont("Helvetica", 10)
        
        for elem in layout_elements:
            if y < 100:
                c.showPage()
                y = height - 50
                
            c.setFont("Helvetica-Bold", 11)
            c.drawString(50, y, f"[{elem.category}] Confidence: {elem.confidence*100:.1f}%")
            y -= 15
            
            c.setFont("Helvetica", 9)
            text_snippet = elem.extracted_text[:120] + "..." if len(elem.extracted_text) > 120 else elem.extracted_text
            c.drawString(50, y, text_snippet)
            y -= 25
            
        c.save()
        return output_path
